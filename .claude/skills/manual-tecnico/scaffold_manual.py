#!/usr/bin/env python3
"""Crea un Manual Tecnico nuevo (esqueleto) a partir de template-base.docx.

No lo uses si el .docx de la HU/bug YA existe en anexos/ -- en ese caso el
skill actualiza el existente con python-docx en vez de llamar este script,
salvo que se este migrando un manual viejo al formato nuevo (ver SKILL.md
seccion "Migracion de manuales existentes").

Formato de las secciones de BACKEND (por defecto, desde 2026-08-03): sigue el
formato de especificacion de API tipo contrato (Informacion General del
Servicio, Descripcion, Especificacion de Request/Response, Catalogo de
mensajes funcionales, Codigos HTTP) -- el mismo usado en los manuales de
HU635 y HU864. Las secciones de FRONTEND (evidencia paso a paso) mantienen
el formato anterior, porque no hay un contrato REST propio que documentar.

Uso:
  python3 scaffold_manual.py \
    --proyecto 416 --modulo SIN --hu HU762 \
    --titulo "Integracion con SARLAFT" \
    --tipo "frontend,backend" \
    --out "/ruta/anexos/Manual Tecnico - 416_SIN_HU762_Integracion SARLAFT.docx"
"""
import argparse
import datetime
import os
import shutil
import sys

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches

TEMPLATE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template-base.docx")
LOGO = os.path.join(os.path.dirname(os.path.abspath(__file__)), "logo-positiva.png")


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:color"), "999999")
        borders.append(el)
    tblPr.append(borders)


def update_header(d, titulo, codigo, fecha, elaboro="Hernan Felipe Nieto"):
    """Actualiza el encabezado de pagina (tabla POSITIVA / PROCESO / Version...)
    con los datos de esta HU/bug -- el template trae texto de la ultima HU usada,
    hay que sobreescribirlo siempre.

    Formato oficial (igual a HU635/HU864): tabla de 2 filas x 3 columnas SIN
    merge vertical -- cada celda es independiente:
      fila1: [logo] [PROCESO: {titulo}] [Version/Clasificacion/Fecha/FORMATO]
      fila2: [Aprobo:] [Reviso:] [Elaboro: {nombre}]
    """
    for section in d.sections:
        header = section.header
        if not header.tables:
            continue
        t = header.tables[0]
        rows = t.rows
        if len(rows) >= 2 and len(rows[0].cells) >= 3:
            rows[0].cells[1].text = f"PROCESO:\n{titulo}"
            rows[0].cells[2].text = f"Version: 1\nClasificacion: Publica\nFecha: {fecha}\n\nFORMATO"
            rows[1].cells[0].text = "Aprobo:"
            rows[1].cells[1].text = "Reviso:"
            rows[1].cells[2].text = f"Elaboro: {elaboro}"


# ---------------------------------------------------------------------------
# Secciones BACKEND: formato spec-de-API (HU635/HU864), placeholders (pendiente)
# ---------------------------------------------------------------------------
SECCIONES_BACKEND_API = [
    "1. Informacion General del Servicio",
    "2. Descripcion del Servicio",
    "3.1 Headers Requeridos",
    "3.2 Estructura del Request Body",
    "3.3 Validaciones de Entrada",
    "4.1 Response Exitoso (HTTP 200)",
    "4.2 Response con Error de Negocio (HTTP 404)",
    "4.3 Response con Error de Validacion (HTTP 400)",
    "5. Codigos de Respuesta HTTP",
    "6. Catalogo de Mensajes Funcionales (si el servicio maneja codigos propios; omitir si no aplica)",
    "7. Arquitectura y Flujo (implementacion: capas, clases, decisiones de diseno)",
]

# Placeholder de tabla para "3.3 Validaciones de Entrada" (columnas oficiales HU635)
TABLA_VALIDACIONES_COLS = ["Campo", "Tipo", "Longitud", "Obligatorio", "Validaciones"]

# Placeholder de tabla para "5. Codigos de Respuesta HTTP" (columnas oficiales HU635)
TABLA_CODIGOS_HTTP = [
    ("Codigo", "Descripcion", "Escenario"),
    ("200", "OK", "(completar)"),
    ("400", "Bad Request", "(completar)"),
    ("401", "Unauthorized", "(completar)"),
    ("404", "Not Found", "(completar)"),
    ("500", "Internal Server Error", "(completar)"),
    ("503", "Service Unavailable", "(completar)"),
]

# ---------------------------------------------------------------------------
# Secciones FRONTEND: formato anterior sin cambios (evidencia paso a paso)
# ---------------------------------------------------------------------------
SECCIONES_FRONTEND = [
    "Evidencia: pantallas del frontend (paso a paso)",
]

SECCIONES_CIERRE = [
    "Resultados de pruebas unitarias",
    "Notas y discrepancias",
]


def build(args):
    if os.path.exists(args.out):
        print(f"YA EXISTE: {args.out}", file=sys.stderr)
        print("No sobrescribas con este script -- abre el .docx existente con "
              "python-docx y actualiza/agrega secciones (o migralo al formato "
              "nuevo si sigue el formato viejo, ver SKILL.md).", file=sys.stderr)
        sys.exit(1)

    shutil.copyfile(TEMPLATE, args.out)
    d = docx.Document(args.out)

    body = d.element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)

    tipos = [t.strip() for t in args.tipo.split(",") if t.strip()]
    es_backend = "backend" in tipos
    es_frontend = "frontend" in tipos

    secciones = []
    if es_backend:
        secciones += SECCIONES_BACKEND_API
    if es_frontend:
        secciones += SECCIONES_FRONTEND
    secciones += SECCIONES_CIERRE

    codigo = f"{args.proyecto}_{args.modulo}_{args.hu}"
    fecha = args.fecha or datetime.date.today().isoformat()

    update_header(d, args.titulo, codigo, fecha)

    # Portada: igual al modelo oficial HU635 -- logo grande centrado, luego
    # "Manual Tecnico" y "{HU} - {titulo}", ambas en negrita y centradas. Sin
    # tabla de metadatos adicional (Historia de Usuario/Repos/Tipo de cambio
    # NO existen en el modelo real; si se necesitan para seguimiento interno,
    # van en Notas y discrepancias, no en la portada).
    for _ in range(6):
        d.add_paragraph("")
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(LOGO, width=Inches(2.2))
    for _ in range(3):
        d.add_paragraph("")

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Manual Tecnico")
    r.bold = True

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{args.hu} - {args.titulo}")
    r.bold = True

    d.add_page_break()
    d.add_paragraph("Contenido")
    for s in secciones:
        d.add_paragraph(s)

    d.add_paragraph("")
    for s in secciones:
        d.add_paragraph(s)
        if s == "3.3 Validaciones de Entrada":
            tv = d.add_table(rows=1, cols=len(TABLA_VALIDACIONES_COLS))
            set_table_borders(tv)
            for i, col in enumerate(TABLA_VALIDACIONES_COLS):
                tv.rows[0].cells[i].text = col
            tv.add_row()
            for cell in tv.rows[1].cells:
                cell.text = "(pendiente)"
        elif s == "5. Codigos de Respuesta HTTP":
            th = d.add_table(rows=len(TABLA_CODIGOS_HTTP), cols=3)
            set_table_borders(th)
            for i, (a, b, c) in enumerate(TABLA_CODIGOS_HTTP):
                th.rows[i].cells[0].text = a
                th.rows[i].cells[1].text = b
                th.rows[i].cells[2].text = c
        else:
            d.add_paragraph("(pendiente)")

    d.save(args.out)
    print(f"Creado: {args.out}")
    print(f"Secciones: {secciones}")


if __name__ == "__main__":
    ap = argparse.ArgumentParser()
    ap.add_argument("--proyecto", default="416")
    ap.add_argument("--modulo", required=True, help="ej. SIN, EMI, INT, REP, COR, FRT")
    ap.add_argument("--hu", required=True, help="ej. HU762 o BUG466519")
    ap.add_argument("--titulo", required=True)
    ap.add_argument("--tipo", required=True, help="frontend,backend o solo uno")
    ap.add_argument("--repos", default="")
    ap.add_argument("--fecha", default="", help="YYYY-MM-DD; por defecto hoy")
    ap.add_argument("--out", required=True)
    build(ap.parse_args())
